"""Loads .docx/.pdf files from sample_docs/ into plain text for chunking."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterator, Union

import pdfplumber
from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

SUPPORTED_EXTENSIONS = {".docx", ".pdf"}


@dataclass
class LoadedDocument:
    """Plain-text body of one source file, plus its filename for provenance.

    Attributes:
        source: Original filename, e.g. "01_PTO_Policy.docx". Gets carried
            through to chunk metadata later so we know where an answer came from.
        text: The whole document as plain text, one line per paragraph/row,
            in reading order.
    """

    source: str
    text: str


def _iter_block_items(document: DocxDocumentType) -> Iterator[Union[Paragraph, Table]]:
    """Walk a docx's paragraphs and tables together, in document order.

    `document.paragraphs` and `document.tables` are two separate, flattened
    lists - reading them separately loses where a table sits relative to the
    surrounding paragraphs (e.g. right after a heading). Walking the
    underlying XML body directly keeps that order intact.

    Args:
        document: A document already opened with python-docx.

    Yields:
        Each top-level Paragraph or Table, in the order it actually appears.
    """
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _table_to_lines(table: Table) -> list[str]:
    """Flatten a docx table into plain text, one "cell | cell | cell" line per row.

    Args:
        table: A docx table.

    Returns:
        One line per row that has at least one non-empty cell. Fully blank
        rows are skipped.
    """
    lines = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    return lines


def load_docx(path: Path) -> LoadedDocument:
    """Load a .docx file, pulling in both paragraphs and tables.

    Reading `document.paragraphs` alone would silently drop any tables -
    which is where a couple of these policies keep their actual numbers
    (accrual tiers, vesting schedules, etc.), so we walk the document in
    order instead and convert tables to text as we go.

    Args:
        path: Path to a .docx file.

    Returns:
        A LoadedDocument holding the file's text.
    """
    document = DocxDocument(str(path))
    lines: list[str] = []
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                lines.append(text)
        else:
            lines.extend(_table_to_lines(block))
    return LoadedDocument(source=path.name, text="\n".join(lines))


def load_pdf(path: Path) -> LoadedDocument:
    """Load a .pdf file, page by page, via pdfplumber's text extraction.

    Args:
        path: Path to a .pdf file.

    Returns:
        A LoadedDocument holding the file's text.
    """
    lines: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            lines.extend(line.strip() for line in page_text.splitlines() if line.strip())
    return LoadedDocument(source=path.name, text="\n".join(lines))


def load_document(path: Path) -> LoadedDocument:
    """Load one file, picking the loader based on its extension.

    Args:
        path: Path to a .docx or .pdf file.

    Returns:
        A LoadedDocument holding the file's text.

    Raises:
        ValueError: If the file isn't a .docx or .pdf.
    """
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return load_docx(path)
    if suffix == ".pdf":
        return load_pdf(path)
    raise ValueError(f"Unsupported file extension '{suffix}' for {path.name}")


def load_documents(directory: Path) -> list[LoadedDocument]:
    """Load every .docx/.pdf file sitting directly in a directory.

    Args:
        directory: Folder containing the source files (not searched recursively).

    Returns:
        One LoadedDocument per supported file, sorted by filename.
    """
    paths = sorted(
        path
        for path in directory.iterdir()
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )
    return [load_document(path) for path in paths]
